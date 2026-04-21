# -*- coding: utf-8 -*-
"""输出 Table0 / Row3 (前期工作进展) 单元格中的所有段落，定位插图位置"""
from pathlib import Path
from docx import Document

REPORT = Path(__file__).resolve().parent.parent / "毕业设计中期报告_宋振豫.docx"
doc = Document(str(REPORT))

out = Path(__file__).resolve().parent / "docx_outline.txt"
tbl = doc.tables[0]
cell = tbl.rows[3].cells[0]  # merged row

with out.open("w", encoding="utf-8") as f:
    f.write(f"Cell has {len(cell.paragraphs)} paragraphs:\n\n")
    for i, p in enumerate(cell.paragraphs):
        t = p.text.strip()
        f.write(f"P{i:3d} [{p.style.name}] {t[:200]}\n")

print(f"wrote {out}")

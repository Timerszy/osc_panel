# -*- coding: utf-8 -*-
"""生成软件架构 PNG 并插入到中期报告 P30 (图1.4 位置)"""
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT   = Path(__file__).resolve().parent.parent
REPORT = ROOT / "毕业设计中期报告_宋振豫.docx"
OUTDOCX= ROOT / "毕业设计中期报告_宋振豫_含架构图.docx"
PNG    = ROOT / "tools" / "arch_fig.png"

# ────────────────────────────────────────────────────────
# 1. 用 Pillow 绘制架构 PNG
# ────────────────────────────────────────────────────────
W, H = 1800, 1100
img = Image.new("RGB", (W, H), "white")
d = ImageDraw.Draw(img)

FONT = "C:/Windows/Fonts/msyh.ttc"
FONT_B = "C:/Windows/Fonts/msyhbd.ttc"

def f(sz, bold=False):
    return ImageFont.truetype(FONT_B if bold else FONT, sz)

def rbox(x1, y1, x2, y2, text, fill, border, font, text_color="#000000",
         radius=8):
    """圆角矩形 + 多行居中文字"""
    d.rounded_rectangle([x1, y1, x2, y2], radius=radius, fill=fill,
                         outline=border, width=2)
    # 居中多行文字
    lines = text.split("\n")
    bbox = d.textbbox((0, 0), "Ay", font=font)
    lh = bbox[3] - bbox[1] + 4
    total_h = lh * len(lines)
    cy = (y1 + y2) / 2 - total_h / 2
    for i, line in enumerate(lines):
        bb = d.textbbox((0, 0), line, font=font)
        tw = bb[2] - bb[0]
        cx = (x1 + x2) / 2 - tw / 2
        d.text((cx, cy + i * lh), line, fill=text_color, font=font)

def arrow(x, y1, y2, color="#555555", w=3):
    """竖直双向箭头"""
    d.line([(x, y1), (x, y2)], fill=color, width=w)
    # 上箭头
    d.polygon([(x, y1), (x - 8, y1 + 12), (x + 8, y1 + 12)], fill=color)
    # 下箭头
    d.polygon([(x, y2), (x - 8, y2 - 12), (x + 8, y2 - 12)], fill=color)

# ── 标题 ──
rbox(40, 25, W - 40, 85,
     "示波器面板显示调节控制系统 — 软件分层架构",
     fill="#2F5496", border="#1F3864",
     font=f(30, bold=True), text_color="#FFFFFF", radius=6)

# ── 四层布局 ──
# y 范围 (每层包含一个总框 + 内部子框)
LAYERS = [
    dict(name="应用逻辑层",    label="Application Logic",  y=(110,  285), fill="#FFF2CC", border="#BF8F00", hdr="#FFE699"),
    dict(name="FreeRTOS 任务层", label="Task Layer",         y=(335,  510), fill="#E2EFDA", border="#375623", hdr="#A9D08E"),
    dict(name="外设驱动层",    label="Hardware Driver",    y=(560,  735), fill="#DDEBF7", border="#1F4E79", hdr="#BDD7EE"),
    dict(name="硬件层",        label="Hardware",           y=(785,  960), fill="#FBE5D6", border="#C65911", hdr="#F4B183"),
]

LABEL_W = 170         # 左侧层名宽度
MARG_L  = 50
MARG_R  = 50

for ly in LAYERS:
    y1, y2 = ly["y"]
    # 层名（左侧圆角）
    rbox(MARG_L, y1, MARG_L + LABEL_W, y2,
         f"{ly['name']}\n{ly['label']}",
         fill=ly["hdr"], border=ly["border"], font=f(17, bold=True))
    # 层主框
    rbox(MARG_L + LABEL_W + 15, y1, W - MARG_R, y2,
         "", fill=ly["fill"], border=ly["border"], font=f(14))

# ── 层内子模块 ──
def draw_sub(y1, y2, items, fill, border, font_size=16):
    """在单一层的主框里，平均分布 items (list of str)"""
    n = len(items)
    box_left  = MARG_L + LABEL_W + 30
    box_right = W - MARG_R - 15
    avail = box_right - box_left
    gap = 15
    bw = (avail - gap * (n - 1)) / n
    pad = 18
    for i, t in enumerate(items):
        x1 = box_left + i * (bw + gap)
        x2 = x1 + bw
        rbox(x1, y1 + pad, x2, y2 - pad, t,
             fill="#FFFFFF", border=border, font=f(font_size))

# 层 1：应用逻辑
draw_sub(LAYERS[0]["y"][0], LAYERS[0]["y"][1], [
    "panel_state.h\n全局状态\n(Mutex 保护)",
    "event_queue\n32 事件队列\n(FreeRTOS)",
    "waveform.c\n波形渲染引擎\n(PSRAM 帧缓冲 112KB)",
    "状态机\nIDLE / SCOPE\nTDR / RUN",
], fill="#FFF2CC", border="#BF8F00")

# 层 2：任务
draw_sub(LAYERS[1]["y"][0], LAYERS[1]["y"][1], [
    "event_task\nprio = 10\n按键 / 编码器分发",
    "display_task\nprio = 5\n100ms 刷波形",
    "adc_task\nprio = 6\n~20Hz 采样",
    "heartbeat_task\nprio = 3\n1s 心跳 + 状态",
], fill="#E2EFDA", border="#375623")

# 层 3：驱动
draw_sub(LAYERS[2]["y"][0], LAYERS[2]["y"][1], [
    "encoder.c\nEC11\n正交解码",
    "button.c\nGPIO + ISR\n20ms 软消抖",
    "led.c\nGPIO\n模式指示",
    "lcd_st7789.c\nSPI2 + DMA\n帧缓冲 112KB",
    "adc_sampler.c\nADC oneshot\n12-bit",
    "uart_comm.c\nUART1\n帧协议",
    "wifi_log.c\nSoftAP + TCP\n日志重定向",
], fill="#DDEBF7", border="#1F4E79", font_size=14)

# 层 4：硬件
draw_sub(LAYERS[3]["y"][0], LAYERS[3]["y"][1], [
    "EC11 × 4\nGPIO 1~8",
    "BTN × 4\nGPIO 9~12",
    "LED × 3\nGPIO 13~15\n+ WS2812 off",
    "ST7789 LCD\n240 × 240\nSPI 40MHz",
    "ADC1 输入\nGPIO 4 / 5",
    "UART1\nGPIO 17 / 18",
    "WiFi SoftAP\nSSID: timer",
], fill="#FBE5D6", border="#C65911", font_size=14)

# ── 层间箭头 ──
def layer_arrows(i):
    y1 = LAYERS[i]["y"][1] + 2
    y2 = LAYERS[i + 1]["y"][0] - 2
    # 取 4 个位置均匀分布
    left  = MARG_L + LABEL_W + 80
    right = W - MARG_R - 80
    for k in range(4):
        x = left + (right - left) * k / 3
        arrow(int(x), y1, y2)

for i in range(len(LAYERS) - 1):
    layer_arrows(i)

# ── 底部说明 ──
tip_y = 985
d.text((60, tip_y),
       "■ 箭头双向表示层间通过函数调用 / 中断回调 / 互斥锁共享数据。",
       fill="#555555", font=f(16))
d.text((60, tip_y + 30),
       "■ 所有全局状态读写均通过 FreeRTOS Mutex 保护；事件队列容量 32，采用生产者-消费者模型。",
       fill="#555555", font=f(16))
d.text((60, tip_y + 60),
       "■ LCD 渲染采用双缓冲 (PSRAM 帧缓冲 → 内部 DMA 行缓冲 → SPI) 消除撕裂。",
       fill="#555555", font=f(16))

img.save(PNG, "PNG", optimize=True)
print(f"PNG saved: {PNG}")

# ────────────────────────────────────────────────────────
# 2. 插入到 docx 的 P30 (图1.4 位置)
# ────────────────────────────────────────────────────────
doc = Document(str(REPORT))
cell = doc.tables[0].rows[3].cells[0]
target = cell.paragraphs[30]   # 空段落 (图1.4 占位)

# 清空段落现有 runs
for r in list(target.runs):
    r._element.getparent().remove(r._element)

target.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = target.add_run()
run.add_picture(str(PNG), width=Cm(15))    # 15cm 宽，适合 A4 纵向页

doc.save(str(OUTDOCX))
print(f"Saved new docx: {OUTDOCX}")
